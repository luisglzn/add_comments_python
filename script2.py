from docx import Document
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from datetime import datetime
import re
from lxml import etree
import json

def split_xml_by_elements(xml):
    root = etree.fromstring(xml)
    paragraphs = root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
    return [etree.tostring(p, encoding='unicode') for p in paragraphs]

def build_txt(paragraphs):
    return [re.sub('<.*?>', '', p) for p in paragraphs]

def add_comments_from_json(doc_path, json_path):
    # Abrir el documento una vez
    doc = Document(doc_path)
    
    # Leer el archivo JSON
    with open(json_path, 'r', encoding='utf-8') as file:
        comments_list = json.load(file)
    
    # Añadir todos los comentarios al documento
    for comment_data in comments_list:
        phrase = comment_data['quote']
        comment_text = comment_data['comment']
        author = comment_data['author']
        add_comment_to_phrase(doc, phrase, comment_text, author)
    
    # Guardar el documento modificado
    new_doc_path = doc_path.replace('.docx', '_suggestions.docx')
    doc.save(new_doc_path)
    print(f"Documento modificado guardado como {new_doc_path}")

def add_comment_to_phrase(doc, phrase, comment_text, author="Anonymous"):
    xml = doc.element.xml
    paragraphs = split_xml_by_elements(xml)
    txt = build_txt(paragraphs)
    
    comment_count = 0
    phrase_regex = re.compile(re.escape(phrase), re.IGNORECASE)
    
    # Obtener el último ID de comentario utilizado
    last_comment_id = get_last_comment_id(doc)
    
    for i, paragraph in enumerate(paragraphs):
        paragraph_text = txt[i]
        if phrase_regex.search(paragraph_text):
            modified_paragraph, comment = add_comment_to_paragraph_end(
                paragraph, phrase, comment_text, author, last_comment_id + comment_count + 1
            )
            # Modificar directamente el XML del párrafo
            update_paragraph_xml(doc, i, modified_paragraph)
            add_comment_to_document(doc, comment)
            comment_count += 1

def get_last_comment_id(doc):
    comments_part = doc.part.comments_part
    if comments_part is None:
        return 0
    comments = comments_part._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment')
    if not comments:
        return 0
    last_comment = comments[-1]
    return int(last_comment.get(qn('w:id')))

def add_comment_to_paragraph_end(paragraph, phrase, comment_text, author, comment_id):
    root = etree.fromstring(paragraph)
    first_occurrence = True

    # Buscar el texto dentro del párrafo
    for elem in root.iter():
        if elem.tag.endswith('t') and elem.text and phrase in elem.text:
            # Dividir el texto en partes: antes, el texto encontrado, y después
            parts = elem.text.split(phrase, 1)
            before_text = parts[0]
            after_text = parts[1] if len(parts) > 1 else ""

            # Ajustar el texto del nodo actual con la parte antes del texto encontrado
            elem.text = before_text

            # Crear un nuevo elemento para el texto encontrado
            found_text = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            text_elem = etree.SubElement(found_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            text_elem.text = phrase

            elem.addnext(found_text)

            # Añadir el espacio inmediatamente al nodo del texto encontrado
            if after_text and after_text.startswith(" "):
                after_text = after_text[1:] # Eliminar el espacio del texto restante

            # Crear un nuevo elemento para el texto restante (si hay)
            if after_text:
                after_text_elem = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                after_text_t = etree.SubElement(after_text_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                after_text_t.text = u'\u00A0' + after_text  # Añadir un espacio de no separación antes del texto restante
                found_text.addnext(after_text_elem)

                # Ajustar el espaciado de caracteres
                rPr = etree.SubElement(after_text_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                spacing = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing')
                spacing.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')

            # Insertar los elementos de comentario alrededor del texto encontrado si es la primera ocurrencia
            if first_occurrence:
                comment_range_start = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeStart')
                comment_range_start.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                found_text.addprevious(comment_range_start)

                comment_range_end = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeEnd')
                comment_range_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                found_text.addnext(comment_range_end)

                comment_reference = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentReference')
                comment_reference.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                comment_range_end.addnext(comment_reference)

                first_occurrence = False
            break # Salir del bucle después de añadir el comentario

    comment = create_comment(str(comment_id), author, comment_text)
    return etree.tostring(root, encoding='unicode'), comment

def update_paragraph_xml(doc, index, modified_paragraph):
    # Actualizar el XML del párrafo en el documento
    body = doc.element.body
    paragraph_elements = body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
    if index < len(paragraph_elements):
        paragraph_elements[index].clear()
        new_paragraph_element = etree.fromstring(modified_paragraph)
        for child in new_paragraph_element:
            paragraph_elements[index].append(child)

def create_comment(comment_id, author, comment_text):
    comment = OxmlElement('w:comment')
    comment.set(qn('w:id'), comment_id)
    comment.set(qn('w:author'), author)
    comment.set(qn('w:date'), datetime.now().isoformat())
    comment.set(qn('w:initials'), ''.join([name[0].upper() for name in author.split() if name]))
    
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = comment_text
    r.append(t)
    p.append(r)
    comment.append(p)
    
    return comment

def add_comment_to_document(doc, comment):
    comments_part = doc.part.comments_part
    if comments_part is None:
        comments_part = doc.part.add_comments_part()
    comments_part._element.append(comment)

def highlight_phrase_in_document(doc_path, phrase):
    doc = Document(doc_path)
    for paragraph in doc.paragraphs:
        highlight_phrase_in_paragraph(paragraph, phrase)
    doc.save(doc_path)


def highlight_phrase_in_paragraph(paragraph, phrase):
    paragraph_xml = paragraph._element
    for elem in paragraph_xml.iter():
        if elem.tag.endswith('t') and elem.text and phrase in elem.text:
            parts = elem.text.split(phrase)
            elem.text = parts[0].rstrip()  # Eliminar espacios al final del texto anterior

            for i, part in enumerate(parts[1:]):
                # Crear el elemento before_text
                before_text = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                before_text_elem = etree.SubElement(before_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                
                if parts[i].endswith(" "):
                    before_text_elem.text = u'\u00A0'  # Espacio no separable
                    rPr = etree.SubElement(before_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                    spacing = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing')
                    spacing.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')
                    elem.addnext(before_text)
                    elem = before_text

                found_text = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                
                rPr = etree.SubElement(found_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                highlight = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight')
                highlight.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'yellow')
                
                text_elem = etree.SubElement(found_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                text_elem.text = phrase
                
                elem.addnext(found_text)

                after_text = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                after_text_elem = etree.SubElement(after_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                
                if part.startswith(" "):
                    after_text_elem.text = u'\u00A0' + part[1:]
                else:
                    after_text_elem.text = part

                rPr = etree.SubElement(after_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                spacing = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing')
                spacing.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')

                found_text.addnext(after_text)
                elem = after_text
                
# Ejemplo de uso
if __name__ == "__main__":
    doc_path =r"C:\Users\luisg\OneDrive\Escritorio\Trabajo\Add comments\add_comments_python\EP3567950-B1__seprotec_es.docx"
    new_doc_path = doc_path.replace('.docx', '_suggestions.docx')
    json_path = r"C:\Users\luisg\OneDrive\Escritorio\Trabajo\Add comments\add_comments_python\errors.json"
    add_comments_from_json(doc_path, json_path)
    #highlight_phrase_in_document(new_doc_path, phrase)
    
    