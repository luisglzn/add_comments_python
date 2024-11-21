from docx import Document
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from datetime import datetime,timezone
import re
from lxml import etree
import json
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from docx.opc.constants import RELATIONSHIP_TYPE as RT


CT.WML_COMMENTS_EXTENDED = "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml"
CT.WML_COMMENTS_IDS = "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsIds+xml"
CT.WML_PEOPLE = "application/vnd.openxmlformats-officedocument.wordprocessingml.people+xml"

RT.COMMENTS_EXTENDED = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
RT.COMMENTS_IDS = "http://schemas.microsoft.com/office/2016/09/relationships/commentsIds"
RT.PEOPLE = "http://schemas.microsoft.com/office/2016/09/relationships/people"

CT.WML_COMMENTS_EXTENSIBLE = "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtensible+xml"
RT.COMMENTS_EXTENSIBLE = "http://schemas.microsoft.com/office/2018/08/relationships/commentsExtensible"

def create_element(name, attrs=None, nsmap=None):
    if ":" in name:
        ns, name = name.split(":")
        if nsmap is None or ns not in nsmap:
            nsmap = {ns: f"http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    else:
        if nsmap is None:
            nsmap = {}
    element = etree.Element(etree.QName(nsmap.get(ns, None), name), nsmap=nsmap)
    if attrs:
        for k, v in attrs.items():
            if ":" in k:
                ns, k = k.split(":")
                element.set(etree.QName(nsmap.get(ns, None), k), v)
            else:
                element.set(k, v)
    return element

def create_comments_extensible_part(document):
    main_document_part = document.part
    comments_extensible_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w16cex:commentsExtensible 
    xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" 
    xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex" 
    xmlns:cx1="http://schemas.microsoft.com/office/drawing/2015/9/8/chartex" 
    xmlns:cx2="http://schemas.microsoft.com/office/drawing/2015/10/21/chartex" 
    xmlns:cx3="http://schemas.microsoft.com/office/drawing/2016/5/9/chartex" 
    xmlns:cx4="http://schemas.microsoft.com/office/drawing/2016/5/10/chartex" 
    xmlns:cx5="http://schemas.microsoft.com/office/drawing/2016/5/11/chartex" 
    xmlns:cx6="http://schemas.microsoft.com/office/drawing/2016/5/12/chartex" 
    xmlns:cx7="http://schemas.microsoft.com/office/drawing/2016/5/13/chartex" 
    xmlns:cx8="http://schemas.microsoft.com/office/drawing/2016/5/14/chartex" 
    xmlns:cr="http://schemas.microsoft.com/office/comments/2020/reactions" 
    xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" 
    xmlns:aink="http://schemas.microsoft.com/office/drawing/2016/ink" 
    xmlns:am3d="http://schemas.microsoft.com/office/drawing/2017/model3d" 
    xmlns:o="urn:schemas-microsoft-com:office:office" 
    xmlns:oel="http://schemas.microsoft.com/office/2019/extlst" 
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" 
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" 
    xmlns:v="urn:schemas-microsoft-com:vml" 
    xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" 
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" 
    xmlns:w10="urn:schemas-microsoft-com:office:word" 
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
    xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" 
    xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" 
    xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex" 
    xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid" 
    xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml" 
    xmlns:w16du="http://schemas.microsoft.com/office/word/2023/wordml/word16du" 
    xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash" 
    xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex" 
    xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" 
    xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" 
    xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" 
    xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" 
    mc:Ignorable="w14 w15 w16se w16cid w16 w16cex w16sdtdh cr w16du wp14">
</w16cex:commentsExtensible>'''
    comments_extensible_part = Part(
        PackURI('/word/commentsExtensible.xml'),
        CT.WML_COMMENTS_EXTENSIBLE,
        comments_extensible_xml,
        main_document_part.package
    )
    rel_id = main_document_part.relate_to(comments_extensible_part, RT.COMMENTS_EXTENSIBLE)
    return comments_extensible_part, rel_id

def add_comment_extensible(comments_extensible_part, comment_id, date_utc):
    try:
        root = etree.fromstring(comments_extensible_part.blob.encode('utf-8'))
    except ValueError:
        root = etree.fromstring(comments_extensible_part.blob)
    
    comment_extensible = etree.SubElement(root, '{http://schemas.microsoft.com/office/word/2018/wordml/cex}commentExtensible')
    comment_extensible.set('{http://schemas.microsoft.com/office/word/2018/wordml/cex}durableId', comment_id)
    comment_extensible.set('{http://schemas.microsoft.com/office/word/2018/wordml/cex}dateUtc', date_utc)
    
    new_blob = etree.tostring(root, encoding='UTF-8', xml_declaration=True, standalone=True)
    comments_extensible_part._blob = new_blob

def create_comments_extended_part(document):
    main_document_part = document.part
    comments_extended_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w15:commentsEx 
    xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" 
    xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex" 
    xmlns:cx1="http://schemas.microsoft.com/office/drawing/2015/9/8/chartex" 
    xmlns:cx2="http://schemas.microsoft.com/office/drawing/2015/10/21/chartex" 
    xmlns:cx3="http://schemas.microsoft.com/office/drawing/2016/5/9/chartex" 
    xmlns:cx4="http://schemas.microsoft.com/office/drawing/2016/5/10/chartex" 
    xmlns:cx5="http://schemas.microsoft.com/office/drawing/2016/5/11/chartex" 
    xmlns:cx6="http://schemas.microsoft.com/office/drawing/2016/5/12/chartex" 
    xmlns:cx7="http://schemas.microsoft.com/office/drawing/2016/5/13/chartex" 
    xmlns:cx8="http://schemas.microsoft.com/office/drawing/2016/5/14/chartex" 
    xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" 
    xmlns:aink="http://schemas.microsoft.com/office/drawing/2016/ink" 
    xmlns:am3d="http://schemas.microsoft.com/office/drawing/2017/model3d" 
    xmlns:o="urn:schemas-microsoft-com:office:office" 
    xmlns:oel="http://schemas.microsoft.com/office/2019/extlst" 
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" 
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" 
    xmlns:v="urn:schemas-microsoft-com:vml" 
    xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" 
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" 
    xmlns:w10="urn:schemas-microsoft-com:office:word" 
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
    xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" 
    xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" 
    xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex" 
    xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid" 
    xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml" 
    xmlns:w16du="http://schemas.microsoft.com/office/word/2023/wordml/word16du" 
    xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash" 
    xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex" 
    xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" 
    xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" 
    xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" 
    xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" 
    mc:Ignorable="w14 w15 w16se w16cid w16 w16cex w16sdtdh w16du wp14">
</w15:commentsEx>'''
    comments_extended_part = Part(
        PackURI('/word/commentsExtended.xml'),
        CT.WML_COMMENTS_EXTENDED,
        comments_extended_xml,
        main_document_part.package
    )
    rel_id = main_document_part.relate_to(comments_extended_part, RT.COMMENTS_EXTENDED)
    return comments_extended_part, rel_id


def create_comments_ids_part(document):
    main_document_part = document.part
    comments_ids_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w16cid:commentsIds 
    xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
    xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex"
    xmlns:cx1="http://schemas.microsoft.com/office/drawing/2015/9/8/chartex"
    xmlns:cx2="http://schemas.microsoft.com/office/drawing/2015/10/21/chartex"
    xmlns:cx3="http://schemas.microsoft.com/office/drawing/2016/5/9/chartex"
    xmlns:cx4="http://schemas.microsoft.com/office/drawing/2016/5/10/chartex"
    xmlns:cx5="http://schemas.microsoft.com/office/drawing/2016/5/11/chartex"
    xmlns:cx6="http://schemas.microsoft.com/office/drawing/2016/5/12/chartex"
    xmlns:cx7="http://schemas.microsoft.com/office/drawing/2016/5/13/chartex"
    xmlns:cx8="http://schemas.microsoft.com/office/drawing/2016/5/14/chartex"
    xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
    xmlns:aink="http://schemas.microsoft.com/office/drawing/2016/ink"
    xmlns:am3d="http://schemas.microsoft.com/office/drawing/2017/model3d"
    xmlns:o="urn:schemas-microsoft-com:office:office"
    xmlns:oel="http://schemas.microsoft.com/office/2019/extlst"
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
    xmlns:v="urn:schemas-microsoft-com:vml"
    xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:w10="urn:schemas-microsoft-com:office:word"
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
    xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
    xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex"
    xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid"
    xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml"
    xmlns:w16du="http://schemas.microsoft.com/office/word/2023/wordml/word16du"
    xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash"
    xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex"
    xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
    xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
    xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
    xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    mc:Ignorable="w14 w15 w16se w16cid w16 w16cex w16sdtdh w16du wp14">
</w16cid:commentsIds>'''
    comments_ids_part = Part(
        PackURI('/word/commentsIds.xml'),
        CT.WML_COMMENTS_IDS,
        comments_ids_xml,
        main_document_part.package
    )
    rel_id = main_document_part.relate_to(comments_ids_part, RT.COMMENTS_IDS)
    return comments_ids_part, rel_id

def add_comment_id(comments_ids_part, para_id, durable_id):
    try:
        root = etree.fromstring(comments_ids_part.blob.encode('utf-8'))
    except ValueError:
        root = etree.fromstring(comments_ids_part.blob)
    
    comment_id = etree.SubElement(root, '{http://schemas.microsoft.com/office/word/2016/wordml/cid}commentId')
    comment_id.set('{http://schemas.microsoft.com/office/word/2016/wordml/cid}paraId', para_id)
    comment_id.set('{http://schemas.microsoft.com/office/word/2016/wordml/cid}durableId', durable_id)
    
    comments_ids_part._blob = etree.tostring(root, encoding='UTF-8', xml_declaration=True, standalone=True)

def create_people_part(document):
    main_document_part = document.part
    people_part = Part(
        PackURI('/word/people.xml'),
        CT.WML_PEOPLE,
        '<w:people xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
        main_document_part.package
    )
    rel_id = main_document_part.relate_to(people_part, RT.PEOPLE)
    return people_part, rel_id



def split_xml_by_elements(xml):
    root = etree.fromstring(xml)
    paragraphs = root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
    return [etree.tostring(p, encoding='unicode') for p in paragraphs]

def build_txt(paragraphs):
    return [re.sub('<.*?>', '', p) for p in paragraphs]

def add_comments_from_json(doc_path, json_path):
    doc = Document(doc_path)
    
    # Create additional parts
    create_comments_extended_part(doc)
    comments_ids_part, _ = create_comments_ids_part(doc)
    create_people_part(doc)
    # Create or get the comments extensible part
    try:
        comments_extensible_part = doc.part.package.part_related_by(RT.COMMENTS_EXTENSIBLE)
    except KeyError:
        comments_extensible_part, _ = create_comments_extensible_part(doc)
    
    with open(json_path, 'r', encoding='utf-8') as file:
        comments_list = json.load(file)
    
    last_comment_id = get_last_comment_id(doc)
    
    for index, comment_data in enumerate(comments_list):
        phrase = comment_data['quote']
        comment_text = comment_data['comment']
        author = comment_data['author']
        comment_id = str(last_comment_id + index + 1)
        
        add_comment_to_phrase(doc, phrase, comment_text, author, comment_id)
         # Generar para_id y durable_id (puedes implementar tu propia lógica para esto)
        para_id = f"{int(comment_id):08X}"
        durable_id = para_id
        
        # Añadir el ID del comentario
        add_comment_id(comments_ids_part, para_id, durable_id)
        
        # Add the comment extensible
        date_utc = datetime.now(timezone.utc).isoformat(timespec='seconds') + 'Z'
        add_comment_extensible(comments_extensible_part, comment_id, date_utc)
    
    new_doc_path = doc_path.replace('.docx', '_suggestions.docx')
    doc.save(new_doc_path)
    print(f"Documento modificado guardado como {new_doc_path}")

def add_comment_to_phrase(doc, phrase, comment_text, author, comment_id):
    xml = doc.element.xml
    paragraphs = split_xml_by_elements(xml)
    txt = build_txt(paragraphs)
    
    phrase_regex = re.compile(re.escape(phrase), re.IGNORECASE)
    
    for i, paragraph in enumerate(paragraphs):
        paragraph_text = txt[i]
        if phrase_regex.search(paragraph_text):
            modified_paragraph, comment = add_comment_to_paragraph_end(
                paragraph, phrase, comment_text, author, comment_id
            )
            
            # Modify the paragraph XML directly
            update_paragraph_xml(doc, i, modified_paragraph)
            add_comment_to_document(doc, comment)
            break  # Stop after adding the first comment

def get_last_comment_id(doc):
    comments_part = doc.part.comments_part
    if comments_part is None:
        return 0
    comments = comments_part._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment')
    if not comments:
        return 0
    last_comment = comments[-1]
    return int(last_comment.get(qn('w:id')))

def add_comment_to_paragraph_end(paragraph, phrase, comment_text, author, comment_id):
    root = etree.fromstring(paragraph)
    first_occurrence = True
    nsmap = {'w': "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    
    for elem in root.iter():
        if elem.tag.endswith('t') and elem.text and phrase in elem.text:
            parts = elem.text.split(phrase, 1)
            before_text = parts[0]
            after_text = parts[1] if len(parts) > 1 else ""
            
            elem.text = before_text
            
            found_text = create_element('w:r', attrs={'w:rsidRPr': '001765A7'}, nsmap=nsmap)
            text_elem = create_element('w:t', nsmap=nsmap)
            text_elem.text = phrase
            found_text.append(text_elem)
            elem.addnext(found_text)
            
            if after_text:
                after_text_elem = create_element('w:r', nsmap=nsmap)
                after_text_t = create_element('w:t', nsmap=nsmap)
                after_text_t.text = after_text
                after_text_elem.append(after_text_t)
                found_text.addnext(after_text_elem)
            
            if first_occurrence:
                comment_range_start = create_element('w:commentRangeStart', attrs={'w:id': str(comment_id)}, nsmap=nsmap)
                comment_range_end = create_element('w:commentRangeEnd', attrs={'w:id': str(comment_id)}, nsmap=nsmap)
                
                comment_reference_r = create_element('w:r', attrs={'w:rsidR': '00A76442'}, nsmap=nsmap)
                comment_reference_rPr = create_element('w:rPr', nsmap=nsmap)
                comment_reference_rStyle = create_element('w:rStyle', attrs={'w:val': 'Refdecomentario'}, nsmap=nsmap)
                comment_reference_rPr.append(comment_reference_rStyle)
                comment_reference_r.append(comment_reference_rPr)
                
                comment_reference = create_element('w:commentReference', attrs={'w:id': str(comment_id)}, nsmap=nsmap)
                comment_reference_r.append(comment_reference)
                
                found_text.addprevious(comment_range_start)
                if after_text:
                    after_text_elem.addprevious(comment_range_end)
                    after_text_elem.addprevious(comment_reference_r)
                else:
                    found_text.addnext(comment_range_end)
                    found_text.addnext(comment_reference_r)
                
                first_occurrence = False
            break
    
    comment = create_comment(str(comment_id), author, comment_text)
    return etree.tostring(root, encoding='unicode'), comment

def update_paragraph_xml(doc, index, modified_paragraph):
    # Actualizar el XML del párrafo en el documento
    body = doc.element.body
    paragraph_elements = body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
    if index < len(paragraph_elements):
        paragraph_elements[index].clear()
        new_paragraph_element = etree.fromstring(modified_paragraph)
        for child in new_paragraph_element:
            paragraph_elements[index].append(child)

def create_comment(comment_id, author, comment_text):
    nsmap = {
        'w': "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        'w14': "http://schemas.microsoft.com/office/word/2010/wordml",
        'w15': "http://schemas.microsoft.com/office/word/2012/wordml",
        'w16cex': "http://schemas.microsoft.com/office/word/2018/wordml/cex",
        'w16cid': "http://schemas.microsoft.com/office/word/2016/wordml/cid",
        'w16': "http://schemas.microsoft.com/office/word/2018/wordml",
        'w16du': "http://schemas.microsoft.com/office/word/2023/wordml/word16du",
        'w16sdtdh': "http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash",
        'wp14': "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
        'wpc': "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
        'wpg': "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
        'wpi': "http://schemas.microsoft.com/office/word/2010/wordprocessingInk",
        'wne': "http://schemas.microsoft.com/office/word/2006/wordml",
    }
    
    comment = create_element('w:comment', attrs={'w:id': comment_id, 'w:author': author, 'w:date': datetime.now().isoformat()}, nsmap=nsmap)
    
    p = create_element('w:p')
    pPr = create_element('w:pPr')
    pStyle = create_element('w:pStyle', attrs={'w:val': 'CommentText'})
    pPr.append(pStyle)
    p.append(pPr)
    
    r = create_element('w:r')
    rPr = create_element('w:rPr')
    rStyle = create_element('w:rStyle', attrs={'w:val': 'CommentReference'})
    rPr.append(rStyle)
    r.append(rPr)
    
    annotationRef = create_element('w:annotationRef')
    r.append(annotationRef)
    p.append(r)
    
    r = create_element('w:r')
    t = create_element('w:t')
    t.text = comment_text
    r.append(t)
    p.append(r)
    
    comment.append(p)
    return comment

def add_comment_to_document(doc, comment):
    comments_part = doc.part.comments_part
    if comments_part is None:
        comments_part = doc.part.add_comments_part()
    comments_part._element.append(comment)

def highlight_phrase_in_document(doc_path, phrase):
    doc = Document(doc_path)
    for paragraph in doc.paragraphs:
        highlight_phrase_in_paragraph(paragraph, phrase)
    doc.save(doc_path)


def highlight_phrase_in_paragraph(paragraph, phrase):
    paragraph_xml = paragraph._element
    for elem in paragraph_xml.iter():
        if elem.tag.endswith('t') and elem.text and phrase in elem.text:
            parts = elem.text.split(phrase)
            elem.text = parts[0].rstrip()  # Eliminar espacios al final del texto anterior

            for i, part in enumerate(parts[1:]):
                # Crear el elemento before_text
                before_text = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                before_text_elem = etree.SubElement(before_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                
                if parts[i].endswith(" "):
                    before_text_elem.text = u'\u00A0'  # Espacio no separable
                    rPr = etree.SubElement(before_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                    spacing = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing')
                    spacing.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')
                    elem.addnext(before_text)
                    elem = before_text

                found_text = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                
                rPr = etree.SubElement(found_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                highlight = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight')
                highlight.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'yellow')
                
                text_elem = etree.SubElement(found_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                text_elem.text = phrase
                
                elem.addnext(found_text)

                after_text = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                after_text_elem = etree.SubElement(after_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                
                if part.startswith(" "):
                    after_text_elem.text = u'\u00A0' + part[1:]
                else:
                    after_text_elem.text = part

                rPr = etree.SubElement(after_text, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                spacing = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing')
                spacing.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')

                found_text.addnext(after_text)
                elem = after_text
                
# Ejemplo de uso
if __name__ == "__main__":
    doc_path =r"C:\Users\luisg\OneDrive\Escritorio\Trabajo\Add comments\add_comments_python\EP3567950-B1__seprotec_es.docx"
    new_doc_path = doc_path.replace('.docx', '_suggestions.docx')
    json_path = r"C:\Users\luisg\OneDrive\Escritorio\Trabajo\Add comments\add_comments_python\errors.json"
    add_comments_from_json(doc_path, json_path)
    #highlight_phrase_in_document(new_doc_path, phrase)
    
    